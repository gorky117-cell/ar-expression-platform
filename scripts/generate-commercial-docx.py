import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles
    # Title
    title = doc.add_heading(level=0)
    run_title = title.add_run("WearWave — Commercial AR Apparel Operating System")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(124, 92, 255) # Electric Violet
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Subtitle
    sub = doc.add_paragraph()
    r_sub = sub.add_run("B2B2C Platform & Authenticated Garment Ownership Blueprint")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(80, 80, 100)
    
    # Metadata Block
    meta_table = doc.add_table(rows=1, cols=1)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = meta_table.cell(0, 0)
    set_cell_background(c, "F3F0FF")
    set_cell_margins(c, 120, 120, 180, 180)
    p_meta = c.paragraphs[0]
    r_m = p_meta.add_run("Version: 1.0 (Master Commercial Architecture)  |  Date: August 2026\nPlatform Production URL: https://ar.aiforall.ltd\nScope: Commercial Monetization, Brand Gateway, WW-Key Claim Protocol & Unit Economics")
    r_m.font.name = "Arial"
    r_m.font.size = Pt(9.5)
    r_m.font.color.rgb = RGBColor(60, 40, 120)
    
    doc.add_paragraph() # Spacing
    
    # Helper to add section headers
    def add_sec(h_text):
        h = doc.add_heading(level=1)
        r = h.add_run(h_text)
        r.font.name = "Arial"
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = RGBColor(30, 30, 45)
        return h

    def add_subsec(h_text):
        h = doc.add_heading(level=2)
        r = h.add_run(h_text)
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(124, 92, 255)
        return h

    # Section 1
    add_sec("1. Executive Summary: The 'Shopify + iOS' for Connected Streetwear")
    p1 = doc.add_paragraph()
    p1.add_run("WearWave is the world's first Decentralized Augmented Reality Apparel Operating System. We bridge physical fashion with interactive digital utility.\n\n"
               "Instead of competing merely as a single t-shirt brand, WearWave operates as the central software, authentication, and computer vision infrastructure powering thousands of apparel brands, streetwear labels, and creators globally.")
    
    # 3 Axioms
    p_ax = doc.add_paragraph()
    r = p_ax.add_run("Core Ecosystem Axioms:\n")
    r.bold = True
    doc.add_paragraph("1. Physical Fabric is the Anchor: Brands print high-end physical garments via Direct-to-Film (DTF) on-demand (MOQ = 1).", style='List Bullet')
    doc.add_paragraph("2. Digital Soul is Dynamic: The garment buyer gains exclusive ownership of the shirt's live digital frequency, updating mood, caption, and 3D animations anytime from their phone.", style='List Bullet')
    doc.add_paragraph("3. The World Interacts in Real Space: Anyone pointing a phone camera at the shirt sees the 3D hologram + real-time wearer vibe and can send live reactions (Likes, Waves, Comments).", style='List Bullet')

    # Section 2
    add_sec("2. Tri-Party Ecosystem Architecture")
    p2 = doc.add_paragraph("The WearWave platform connects three distinct stakeholders into a unified digital fashion ecosystem:")
    
    doc.add_paragraph("• Platform Owner (WearWave Master OS): Authorizes brands, compiles precision computer vision descriptors (.mind), issues cryptographic Garment Claim Tokens (WW-Keys), hosts WebAR viewports, and collects royalties & SaaS subscriptions.", style='List Bullet')
    doc.add_paragraph("• Authorized Brand Partners: Streetwear labels, independent designers, and merchandise creators who design drops, print DTF garments on-demand, distribute shirts with WW-Keys, and pay per-unit royalties.", style='List Bullet')
    doc.add_paragraph("• End Consumers (Wearers & Responders): Buyers scan their WW-Key to claim digital ownership, control their daily frequency, and receive live interactive pulses from anyone scanning their shirt in physical space.", style='List Bullet')

    # Section 3
    add_sec("3. Garment Authentication & Ownership Claim Protocol")
    p3 = doc.add_paragraph("To ensure exclusive digital control and prevent counterfeit digital hijacking, every physical garment is bound to a Cryptographic Garment Claim Token (WW-Key).")
    
    add_subsec("A. Manufacturing & Key Generation")
    doc.add_paragraph("When a partner brand initiates a production batch, WearWave automatically generates a batch of unique, single-use activation keys (e.g. WW-8492-BTFL-7X). The key is embedded via an inside neck label scratch-off QR code or a flexible waterproof NFC tag.", style='List Bullet')
    
    add_subsec("B. Buyer Claim & Digital Binding")
    doc.add_paragraph("When the buyer receives the garment, they scan the inside QR/NFC tag. The WearWave platform validates the key, assigns exclusive digital control of that garment's AR channel to the buyer's account, and burns the activation token so it can never be claimed again.", style='List Bullet')

    # Section 4
    add_sec("4. Monetization & Revenue Architecture")
    p4 = doc.add_paragraph("WearWave operates a multi-tiered, high-margin revenue model combining B2B SaaS, Unit Royalties, D2C Flagship Drops, and Premium Wearer Perks:")
    
    # Revenue Table
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Revenue Stream", "Pricing / Model", "Margin & Impact"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_background(cell, "7C5CFF")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    data = [
        ("B2B Per-Unit Royalty", "₹80 – ₹180 INR ($1.00 – $2.50 USD) per shirt", "95%+ Pure Software Margin"),
        ("Brand Enterprise SaaS", "$99 – $499 / month subscription", "Recurring B2B ARR (Analytics & Multi-Target)"),
        ("In-House D2C Flagship Drops", "₹1,499 – ₹2,499 INR ($25 – $45 USD)", "70% – 85% Gross Profit Margin"),
        ("Premium Wearer Perks", "$2.99 / month (Custom particle trails & badges)", "High-LTV Consumer Subscription")
    ]
    
    for row_idx, row_data in enumerate(data, start=1):
        bg = "F9F8FF" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            
    doc.add_paragraph()

    # Section 5
    add_sec("5. Direct-to-Film (DTF) On-Demand Manufacturing Blueprint (MOQ = 1)")
    doc.add_paragraph("• Commercial Machinery: Industrial Epson i3200 Dual/Quad-Head and Mimaki TxF150-75 printers printing at 1440 to 2400 DPI.", style='List Bullet')
    doc.add_paragraph("• Opaque White Barrier: 5-channel CMYK + solid white underbase prevents color bleed into dark cotton, ensuring sharp contrast edges (<0.1mm) for computer vision.", style='List Bullet')
    doc.add_paragraph("• Zero Glare Matte Finish: Diffuse matte texture allows phone sensors to lock onto the artwork from 10 to 15+ feet away and up to 60° angles.", style='List Bullet')
    doc.add_paragraph("• Wash Durability: TPU hot-melt adhesive withstands 50+ machine washes with 100% AR tracking retention.", style='List Bullet')

    # Section 6
    add_sec("6. Technical Safeguards & Proprietary Math")
    doc.add_paragraph("• Zero-App WebAR: Instant launch in mobile Safari/Chrome via standard WebGL/WebRTC without app store friction.", style='List Bullet')
    doc.add_paragraph("• 6-DOF Jitter Smoother (matrix-smoother): Exponential Moving Average (alpha = 0.18) eliminates body movement and fabric flutter jitter.", style='List Bullet')
    doc.add_paragraph("• Distance Auto-Scaler (distance-auto-scaler): Dynamically scales floating story badges up to 2.5x for long-distance legibility.", style='List Bullet')
    doc.add_paragraph("• Permanent Social Memory: Real-time synchronization via Supabase PostgreSQL, permanently storing likes, waves, and comments.", style='List Bullet')

    # Section 7
    add_sec("7. Phased Implementation Roadmap")
    doc.add_paragraph("• Phase 1 (Completed): Markerless WebAR scanner, 1080p stream, dual-target tracking, Supabase reactions, clean 2-card feed.", style='List Bullet')
    doc.add_paragraph("• Phase 2 (In Progress): Dynamic 3D GLTF model selection (Sailing boat, dragon, birds), figure-8 flight paths, room-locking spatial anchor fallback.", style='List Bullet')
    doc.add_paragraph("• Phase 3 (Commercial Rollout): Brand Partner Portal, Cryptographic WW-Key Generator, Garment Claim & Authentication Gateway.", style='List Bullet')
    doc.add_paragraph("• Phase 4 (Global Scale): Multi-brand marketplace, automated royalty billing, and mobile NFC 1-tap claim integration.", style='List Bullet')

    # Save Document
    downloads_path = os.path.join(os.environ.get('USERPROFILE', 'C:\\Users\\lenovo'), 'Downloads', 'WearWave_Commercial_Platform_Blueprint.docx')
    doc.save(downloads_path)
    print(f"SUCCESS: Saved Word Document to {downloads_path}")

if __name__ == "__main__":
    create_document()
